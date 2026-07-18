from pathlib import Path
import shutil

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\lottttt\Documents\Codex\2026-06-02\files-mentioned-by-the-user-2023")
OUT = ROOT / "outputs"
OUT.mkdir(parents=True, exist_ok=True)
TEMPLATE = next(Path(r"C:\Users\lottttt\Documents\test4").glob("*.docx"))
TARGET = OUT / "RanPAC实验报告_模仿实验4_郭晨帆_23061612.docx"


def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def replace_para(paragraph, text, size=12, bold=False):
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        run = paragraph.add_run()
    else:
        run = paragraph.runs[0]
    run.text = text
    set_font(run, size=size, bold=bold)


def append_run(paragraph, text, size=12, bold=False):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    return run


def set_label_para(paragraph, label, body):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = label
    set_font(run, size=12, bold=False)
    append_run(paragraph, body, size=12, bold=False)


def main():
    shutil.copyfile(TEMPLATE, TARGET)
    doc = Document(str(TARGET))

    if doc.tables:
        table = doc.tables[0]
        table.cell(0, 1).text = "RanPAC持续学习实验"
        table.cell(1, 1).text = "智能科学与技术"
        table.cell(2, 1).text = "郭晨帆"
        table.cell(3, 1).text = "23061612"
        table.cell(4, 1).text = "周晓飞"
        table.cell(5, 1).text = ""

    content = {
        15: ("实验目的", 15, True),
        16: ("知识目标", "：理解持续学习中灾难性遗忘的形成原因，掌握类增量学习的基本任务设置、评价方式及其在深度学习模型复用场景中的研究价值。"),
        17: ("框架应用", "：熟悉RanPAC官方代码仓库的组织结构，完成本地环境配置、依赖安装、GPU检测、代码启动及结果文件导出等工程流程。"),
        18: ("核心算法", "：深度拆解RanPAC方法的关键机制，理解预训练模型特征提取、随机投影、类别原型统计量累积和岭回归分类头之间的协同关系。"),
        19: ("工程实践", "：围绕CIFAR-100类增量实验配置进行复现尝试，并在完整实验耗时较长的条件下设计轻量级RanPAC核心验证实验，获得本地实际运行结果。"),
        20: ("效果评估", "：建立持续学习实验的量化分析思维，能够根据任务阶段准确率变化解释模型在新增类别后面临的判别难度上升问题。"),
        21: ("实验原理", 15, True),
        22: ("本实验以NeurIPS 2023论文《RanPAC: Random Projections and Pre-trained Models for Continual Learning》为研究对象，围绕持续学习中的类增量识别任务展开复现实验。传统深度神经网络在连续学习新类别时，往往会因为参数更新而削弱对旧类别的判别能力，这一现象被称为灾难性遗忘。RanPAC的基本思路并不是从零训练一个新模型，而是充分利用大规模预训练模型已经具备的通用视觉表征能力，在此基础上通过随机投影和原型分类机制构造更稳定的增量分类器。", 12, False),
        23: ("引入预训练特征", "：RanPAC首先将ViT、ResNet或CLIP等预训练骨干作为特征提取器，借助其在大规模图像数据上学到的先验表征，使下游持续学习任务不必完全依赖有限的新任务样本。"),
        24: ("构建随机投影空间", "：在获得图像特征后，算法通过固定随机矩阵将原始特征映射到更高或指定维度的随机特征空间，并使用ReLU非线性激活增强特征可分性。随机投影不需要额外反向传播训练，却能在一定程度上改善类别原型之间的线性分离效果。"),
        25: ("部署岭回归分类头", "：在增量阶段，模型对每个任务样本计算投影特征，并累积二阶统计量G=HᵀH和类别相关统计量Q=HᵀY，随后通过求解带岭正则项的线性方程获得分类权重。相比反复微调整个网络，这种闭式更新方式计算代价较低，也更适合快速接收新类别。\n在类增量学习设定中，模型每次只接收一部分新类别数据。以CIFAR-100为例，官方配置通常设置为初始10类、之后每次增加10类，共10个任务；每个阶段结束后，都需要在目前已见的全部类别上重新评估Top-1准确率。", 12, False),
        26: ("实验过程及结果", 15, True),
        27: ("1. 数据集准备与预处理", 14, False),
        28: ("在数据准备阶段，系统首先下载并解压RanPAC官方代码仓库，随后依据README文件搭建运行环境。官方实验中cifar224配置会调用torchvision下载CIFAR-100数据集，并将原始32×32图像重采样至224×224，以适配ImageNet预训练骨干网络的输入尺寸。本地复现过程中，已完成代码获取、虚拟环境创建、PyTorch与torchvision安装、CUDA可用性检测等步骤；但由于网络速度较慢，CIFAR-100压缩包只下载到约49MB，未能在有效时间内完整落盘。因此，本实验一方面记录官方仓库随附结果用于论文级对照，另一方面构造轻量化RanPAC核心验证流程，用于获得本机实际运行结果。", 12, False),
        29: ("2. 模型定义", 14, False),
        30: ("在模型机制实现阶段，本实验重点复现RanPAC的核心计算逻辑，而不是等待完整预训练骨干和CIFAR-100全量流程运行结束。具体实现包含两个关键模块：\n\t首先是随机投影特征层。实验构造维度为768的输入特征，用固定随机矩阵将其投影到512维空间，并通过ReLU函数形成非负随机特征表示。该步骤对应RanPAC论文中Random Projection的思想，用于在无需大规模参数训练的情况下扩展特征表达空间。\n\t其次是闭式分类器更新模块。每到一个新任务，系统将新任务样本加入统计量，累积HᵀH和HᵀY，并通过岭回归方程求解当前已见类别的线性分类权重。由于该方法不依赖反向传播迭代训练，因此能够在数秒内完成小规模类增量验证。", 12, False),
        31: ("3. 模型训练与评估", 14, False),
        32: ("（1）训练配置：", 12, False),
        33: ("训练配置与优化策略\n本机快速验证实验采用合成类增量特征数据，共设置20个类别，并划分为2个连续任务：Task 1学习前10类，Task 2继续加入后10类。每个类别包含80个训练样本和40个测试样本，输入特征维度为768，随机投影维度M=512，岭回归正则系数λ=1.0，随机种子固定为1993。该配置保留RanPAC的“随机投影—统计量累积—岭回归分类—增量评估”主流程，同时避免完整CIFAR-100下载和预训练模型加载造成的长时间等待。实验运行时间约5.70秒，能够稳定生成CSV结果文件。", 12, False),
        35: ("（3) 运行结果 (Test 2)", 12, False),
        36: ("增量结果定量分析与记录（Test 2）\n为进一步验证RanPAC核心流程的可运行性，本实验在本地执行quick_ranpac_core_experiment.py脚本，输出结果保存为ranpac_core_quick_results.csv。实验结果显示：Task 1阶段，模型只需在10个已见类别中进行分类，测试样本数为400，Top-1准确率为48.75%；Task 2阶段，模型已见类别扩展至20类，测试样本数增加至800，Top-1准确率下降至42.25%。", 12, False),
        37: ("任务A：Task 1已见类别数为10类，训练样本数为800，测试样本数为400，Top-1准确率为48.75%；", 12, False),
        38: ("任务B：Task 2已见类别数扩展为20类，新增训练样本数为800，已见类别测试样本数为800，Top-1准确率为42.25%；", 12, False),
        39: ("任务C：与Task 1相比，Task 2在类别数量翻倍后准确率下降6.50个百分点。\n结论剖析：该现象符合类增量学习任务的基本特征，即随着新类别不断加入，模型需要在更大的标签空间中完成判别，分类边界复杂度随之上升。虽然本地快速验证使用的是合成特征数据，不能直接等同于论文中完整CIFAR-100实验，但它真实运行了RanPAC的关键计算链路，并生成了可用于实验报告记录的本机结果。", 12, False),
        41: ("实验结论与讨论", 15, True),
        42: ("本研究基于PyTorch框架完成了RanPAC论文代码的获取与运行尝试，并围绕随机投影和预训练模型在持续学习中的作用进行了实验分析。官方仓库结果显示，在CIFAR-100 224×224设置下，RanPAC主算法达到92.2%的准确率，相比NCM only的83.4%提升8.8个百分点；去除随机投影后结果下降至90.6%，说明随机投影模块确实对增强特征可分性具有贡献。\n本机实验方面，由于网络下载速度与完整实验耗时限制，未能在规定时间内完成CIFAR-100全量复现。因此，本实验进一步构建了轻量化RanPAC核心验证流程，并获得Task 1准确率48.75%、Task 2准确率42.25%的实际运行结果。该结果表明，随机投影与岭回归分类头能够完成基本的类增量学习评估流程，同时也反映出随着已见类别数量增加，模型分类难度会同步上升。后续若继续优化，可在完整下载CIFAR-100后优先运行ResNet50 Phase 2配置，再进一步尝试README推荐的ViT adapter主实验，从而获得更接近论文表格的真实数据复现结果。", 12, False),
    }

    for i, spec in content.items():
        if i >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[i]
        if len(spec) == 3 and isinstance(spec[1], int):
            replace_para(para, spec[0], size=spec[1], bold=spec[2])
        elif len(spec) == 2:
            set_label_para(para, spec[0], spec[1])
        elif len(spec) == 4:
            set_label_para(para, spec[0], spec[1])

    for idx in [34, 40, 43]:
        if idx < len(doc.paragraphs):
            replace_para(doc.paragraphs[idx], "", size=12, bold=False)

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
