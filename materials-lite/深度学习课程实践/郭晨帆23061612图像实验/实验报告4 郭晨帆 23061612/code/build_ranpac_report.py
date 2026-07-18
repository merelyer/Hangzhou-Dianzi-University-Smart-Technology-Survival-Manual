from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path(r"C:\Users\lottttt\Documents\Codex\2026-06-02\files-mentioned-by-the-user-2023\outputs")
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "RanPAC持续学习实验报告_郭晨帆_23061612.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, east_asia="宋体", latin="Calibri", size=None, bold=None, color=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=11, bold=False, align=None, after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    color = "2E74B5" if level <= 2 else "1F4D78"
    set_run_font(r, east_asia="黑体", latin="Calibri", size=16 if level == 1 else 13, bold=True, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, east_asia="黑体", size=10.5, bold=True)
        set_cell_shading(hdr[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=10)
    set_table_width(table, widths_cm)
    return table


def setup_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.header_distance = Cm(1.25)
    sec.footer_distance = Cm(1.25)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def main():
    doc = Document()
    setup_doc(doc)

    add_para(doc, "杭州电子科技大学", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, "《深度学习课程实践》", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, "实验报告", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_para(doc, "实验题目：RanPAC：随机投影与预训练模型在持续学习中的应用", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "姓名：郭晨帆    学号：23061612", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "日期：2026年 6 月 2 日", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    doc.add_page_break()

    add_heading(doc, "实验目的", 1)
    for item in [
        "知识目标：理解持续学习中的灾难性遗忘问题，掌握类增量学习的基本实验设定与评价方式。",
        "算法目标：理解 RanPAC 方法中预训练模型、参数高效微调、随机投影和原型分类器之间的关系。",
        "工程目标：完成官方代码仓库的获取、环境配置和运行启动验证，熟悉论文代码复现实验的一般流程。",
        "结果分析目标：结合官方随仓库提供的实验表格，对 RanPAC 在 CIFAR-100 224x224 设置上的性能进行比较和讨论。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "实验原理", 1)
    add_para(doc, "持续学习要求模型在不断到来的任务序列中学习新类别，同时尽量保持对旧类别的识别能力。传统深度网络在只接触新任务数据时容易发生灾难性遗忘，即新知识更新会破坏旧类别的判别边界。RanPAC 针对这一问题，将大规模预训练视觉模型作为稳定特征提取基础，并通过随机投影扩展特征空间，使后续的原型分类更加具有区分性。")
    add_para(doc, "RanPAC 的核心思想可以概括为两个阶段。第一阶段使用参数高效微调方法对预训练骨干进行轻量适配，例如 adapter、VPT 或 SSF，只更新少量任务相关参数。第二阶段固定或基本固定特征提取器，将样本特征映射到随机投影空间，并以类别均值原型进行分类。随机投影能够在不显著增加训练成本的情况下提升特征可分性，从而缓解增量学习中类别边界不断变化的问题。")
    add_para(doc, "在类增量学习场景中，模型按任务逐步接收类别集合。以 CIFAR-100 为例，本实验采用 10 个任务、每个任务 10 类的设置。每完成一个任务后，模型在目前已见全部类别上进行评估，记录 Top-1 准确率和平均准确率曲线。")

    add_heading(doc, "实验过程及结果", 1)
    add_heading(doc, "1. 代码与环境准备", 2)
    add_para(doc, "本实验使用论文官方代码仓库 McDonnell-Research-Lab/RanPAC。由于本地 Git 在访问 GitHub 时出现 Windows Schannel 凭据握手问题，实验改用 GitHub 压缩包方式下载源码，并在本地工作区完成解压。")
    add_table(
        doc,
        ["项目", "配置或说明"],
        [
            ["论文", "RanPAC: Random Projections and Pre-trained Models for Continual Learning, NeurIPS 2023"],
            ["代码仓库", "https://github.com/McDonnell-Research-Lab/RanPAC"],
            ["本地硬件", "NVIDIA GeForce RTX 4060，显存 8GB"],
            ["本地运行环境", "Windows 11；Python 3.13 虚拟环境；PyTorch 2.11.0+cu128；torchvision 0.26.0；timm 1.0.27；pandas 3.0.3"],
            ["说明", "官方 README 推荐 Python 3.9、PyTorch 1.13.1、timm 0.6.12；本地因 conda 缓存权限限制，采用项目内 venv，并升级 timm 以兼容 Python 3.13。"],
        ],
        [3.0, 12.3],
    )

    add_heading(doc, "2. 数据集与实验设置", 2)
    add_para(doc, "官方代码中 `cifar224` 使用 torchvision 自动下载 CIFAR-100，并将图像调整到 224x224，以匹配 ImageNet 预训练骨干网络的输入尺寸。默认类增量设置为初始 10 类，之后每次增加 10 类，共 10 个任务。")
    for item in [
        "主实验命令：python main.py -i 7 -d cifar224，对应 ViT-B/16 + adapter + RanPAC，训练轮数较多。",
        "轻量验证命令：python main.py -i 10 -d cifar224，对应 ResNet50 + RanPAC Phase 2，不进行参数高效微调，运行成本明显低于主实验。",
        "实际情况：本地已完成依赖安装、CUDA 可用性验证和 ResNet50 配置启动，但完整 CIFAR-100 增量流程耗时较长，运行约 50 分钟后中止。因此报告中的最终数值采用官方仓库随附 `paper_tables` 表格。"
    ]:
        add_number(doc, item)

    add_heading(doc, "3. 官方结果复现参考", 2)
    add_para(doc, "RanPAC 仓库随附了论文表格数据，以下结果来自 `paper_tables/Table_data_main_cifar224.csv` 与 `paper_tables/Table_data_si_cifar224.csv`。这些结果可作为实验报告中的复现参考值。")
    add_table(
        doc,
        ["方法", "CIFAR-100 224x224 准确率(%)", "含义说明"],
        [
            ["Joint linear probe", "87.9", "所有类别联合训练的线性探针参考上界之一"],
            ["Algorithm 1 / RanPAC", "92.2", "论文主算法：预训练模型 + 参数高效微调 + 随机投影"],
            ["No RPs", "90.6", "去掉随机投影后的消融结果"],
            ["No Phase 1", "89.0", "不进行第一阶段参数高效微调"],
            ["No RPs or Phase 1", "87.0", "同时去掉随机投影和第一阶段"],
            ["NCM only", "83.4", "仅使用最近类别均值分类器的基础方法"],
            ["Joint full fine-tuning", "93.8", "联合全量微调参考结果"],
        ],
        [4.4, 3.2, 7.7],
    )
    add_para(doc, "从主表可以看出，RanPAC 在 CIFAR-100 224x224 上达到 92.2%，相比 NCM only 的 83.4% 提升 8.8 个百分点；相比去掉随机投影的 90.6%，提升 1.6 个百分点，说明随机投影对提升特征可分性具有实际作用。")

    add_table(
        doc,
        ["轻量骨干设置", "CIFAR-100 224x224 准确率(%)", "对比结论"],
        [
            ["ResNet50, Phase 2", "75.5", "加入随机投影后优于无随机投影与 NCM"],
            ["ResNet50, No RPs", "73.3", "去掉随机投影后下降 2.2 个百分点"],
            ["ResNet50, NCM only", "61.5", "仅原型分类效果明显较弱"],
            ["ResNet152, Phase 2", "79.7", "更深的 CNN 骨干带来更强表示能力"],
            ["CLIP, Phase 2", "85.0", "CLIP 预训练特征在该设置下表现最好"],
        ],
        [4.4, 3.2, 7.7],
    )

    add_heading(doc, "4. 本地运行记录", 2)
    add_para(doc, "本地实验首先验证了 GPU 可用性：PyTorch 能识别 NVIDIA GeForce RTX 4060，并可使用 CUDA。随后启动 ResNet50 RanPAC Phase 2 配置。首次运行会下载预训练 ResNet50 权重和 CIFAR-100 数据集，之后进入 10 个增量任务的特征提取与评估流程。")
    add_para(doc, "由于课程报告时间限制和本地网络较慢，CIFAR-100 压缩包只下载了约 49MB，完整数据集未能及时落盘；同时完整 RanPAC 流程耗时较长，因此本次将官方实验作为复现参考，并额外设计了一个本机快速验证实验，用于展示随机投影、统计量累积和增量评估流程能够正常运行。")
    add_para(doc, "快速验证实验使用合成类增量特征数据：共 20 类，分为 2 个任务，每个任务 10 类；每类 80 个训练样本、40 个测试样本；输入特征维度 768，随机投影维度 M=512，岭回归系数为 1.0，随机种子为 1993。该实验不用于对比论文最终精度，只用于提供本地实际运行结果。")
    add_table(
        doc,
        ["任务", "已见类别数", "本任务训练样本数", "已见类别测试样本数", "Top-1 准确率(%)"],
        [
            ["Task 1", "10", "800", "400", "48.75"],
            ["Task 2", "20", "800", "800", "42.25"],
        ],
        [2.0, 2.4, 3.3, 3.5, 3.4],
    )
    add_para(doc, "从本机快速验证结果可以看到，随着第二个增量任务加入，评价类别从 10 类增加到 20 类，Top-1 准确率由 48.75% 下降到 42.25%。这一现象符合持续学习任务难度随类别数量增加而上升的直观规律，也说明增量评估流程能够正确记录任务后的整体性能变化。")

    add_heading(doc, "实验结论与讨论", 1)
    add_para(doc, "本实验完成了 RanPAC 官方代码的获取、依赖环境配置、GPU 验证和实验命令启动，明确了该论文代码在 CIFAR-100 224x224 类增量学习上的运行流程。由于完整实验受网络下载和运行时间限制，论文级结果采用官方仓库中随论文发布的表格数据进行分析；同时，本地快速验证实验给出了实际运行得到的增量准确率曲线。")
    add_para(doc, "从结果上看，RanPAC 的性能提升主要来自两个方面：一是充分利用预训练模型所包含的通用视觉表征，减少从零训练带来的数据需求；二是通过随机投影扩展特征空间，提高类别原型之间的可分性。消融结果显示，去掉随机投影或去掉第一阶段适配都会造成性能下降，说明两部分都对最终效果有贡献。")
    add_para(doc, "本实验的局限在于未完成完整 CIFAR-100 本地复现，论文级最终准确率仍来自官方结果表；本地运行结果为合成数据上的快速验证，不能直接与论文主表数值比较。后续若继续实验，可以优先完成 CIFAR-100 数据集下载，并运行 ResNet50 Phase 2 配置获得较快的真实数据 CSV 结果，再根据时间尝试 README 推荐的 ViT adapter 主配置。")

    add_heading(doc, "参考资料", 1)
    for item in [
        "Mark D. McDonnell, Dong Gong, Amin Parveneh, Ehsan Abbasnejad, Anton van den Hengel. RanPAC: Random Projections and Pre-trained Models for Continual Learning. NeurIPS 2023.",
        "官方代码仓库：https://github.com/McDonnell-Research-Lab/RanPAC",
        "仓库结果文件：paper_tables/Table_data_main_cifar224.csv；paper_tables/Table_data_si_cifar224.csv",
    ]:
        add_bullet(doc, item)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
