from pathlib import Path
import shutil
import sys

import pandas as pd


ROOT = Path(r"C:\Users\lottttt\Documents\Codex\2026-06-02\files-mentioned-by-the-user-2023")
OUT = ROOT / "outputs"
REPO = ROOT / "work" / "RanPAC-main"
REPORT_IN = OUT / "RanPAC实验报告_模仿实验4_郭晨帆_23061612.docx"
REPORT_OUT = OUT / "RanPAC实验报告_带图表_郭晨帆_23061612.docx"
CHART1 = OUT / "ranpac_local_accuracy_curve.png"
CHART2 = OUT / "ranpac_official_cifar224_comparison.png"


def font_run(run, size=12, bold=False):
    from docx.oxml.ns import qn
    from docx.shared import Pt
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def make_charts():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    quick = pd.read_csv(OUT / "ranpac_core_quick_results.csv")
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(6.6, 3.8), dpi=200)
    ax.plot(quick["task"], quick["top1_accuracy"], marker="o", linewidth=2.2, color="#2E74B5")
    ax.bar(quick["task"], quick["top1_accuracy"], alpha=0.18, color="#2E74B5", width=0.5)
    for x, y in zip(quick["task"], quick["top1_accuracy"]):
        ax.text(x, y + 1.0, f"{y:.2f}%", ha="center", fontsize=9)
    ax.set_title("本机快速验证实验：增量任务Top-1准确率")
    ax.set_xlabel("增量任务")
    ax.set_ylabel("Top-1准确率 (%)")
    ax.set_xticks(quick["task"])
    ax.set_ylim(0, max(60, quick["top1_accuracy"].max() + 10))
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    fig.tight_layout()
    fig.savefig(CHART1, bbox_inches="tight")
    plt.close(fig)

    official = pd.read_csv(REPO / "paper_tables" / "Table_data_main_cifar224.csv")
    official = official[official["Unnamed: 0"] != "runtime"].copy()
    official = official.rename(columns={"Unnamed: 0": "method", "cifar224": "accuracy"})
    keep = ["NCM only", "No RPs or Phase 1", "No Phase 1", "No RPs", "Algorithm 1", "Joint full fine-tuning"]
    official = official[official["method"].isin(keep)]
    official["method"] = pd.Categorical(official["method"], categories=keep, ordered=True)
    official = official.sort_values("method")

    fig, ax = plt.subplots(figsize=(7.2, 4.2), dpi=200)
    colors = ["#A7A9AC", "#9CB9D8", "#7EA5CC", "#5C8DBB", "#2E74B5", "#1F4D78"]
    bars = ax.barh(official["method"], official["accuracy"], color=colors)
    for bar, val in zip(bars, official["accuracy"]):
        ax.text(val + 0.4, bar.get_y() + bar.get_height() / 2, f"{val:.1f}%", va="center", fontsize=9)
    ax.set_title("官方CIFAR-100 224x224结果对比")
    ax.set_xlabel("准确率 (%)")
    ax.set_xlim(80, 96)
    ax.grid(axis="x", linestyle="--", alpha=0.3)
    fig.tight_layout()
    fig.savefig(CHART2, bbox_inches="tight")
    plt.close(fig)


def add_chart_after_paragraph(doc, marker, image_path, caption):
    from docx.shared import Cm
    for i, para in enumerate(doc.paragraphs):
        if marker in para.text:
            p = doc.paragraphs[i + 1]._p.addnext
            # python-docx has no direct insert-after API; use a temporary paragraph at end
            break
    else:
        i = len(doc.paragraphs) - 1

    new_p = doc.add_paragraph()
    new_p.alignment = 1
    run = new_p.add_run()
    run.add_picture(str(image_path), width=Cm(13.5))
    cap = doc.add_paragraph()
    cap.alignment = 1
    r = cap.add_run(caption)
    font_run(r, size=10, bold=False)

    # Move the two newly appended paragraphs after the marker paragraph.
    body = doc._body._element
    pic_p = new_p._p
    cap_p = cap._p
    body.remove(pic_p)
    body.remove(cap_p)
    anchor = doc.paragraphs[i]._p
    anchor.addnext(cap_p)
    anchor.addnext(pic_p)


def main():
    mode = sys.argv[1] if len(sys.argv) > 1 else "all"
    if mode in ("all", "charts"):
        make_charts()
    if mode == "charts":
        print(CHART1)
        print(CHART2)
        return
    from docx import Document
    from docx.shared import Cm
    shutil.copyfile(REPORT_IN, REPORT_OUT)
    doc = Document(str(REPORT_OUT))
    add_chart_after_paragraph(doc, "Task 2已见类别数扩展为20类", CHART1, "图1 本机RanPAC核心验证实验的增量准确率变化")
    add_chart_after_paragraph(doc, "官方仓库结果显示", CHART2, "图2 RanPAC官方CIFAR-100 224x224主要方法准确率对比")
    doc.save(REPORT_OUT)
    print(REPORT_OUT)
    print(CHART1)
    print(CHART2)


if __name__ == "__main__":
    main()
